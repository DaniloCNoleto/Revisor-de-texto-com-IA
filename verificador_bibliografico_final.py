# ğŸ“Œ Revisor BibliogrÃ¡fico Dossel com filtro por categoria, Ã­ndice absoluto e seleÃ§Ã£o de pastas
import os
import sys
import time
import openai
import re
import json
import traceback
import requests
import difflib
from docx import Document
from openpyxl import Workbook, load_workbook
from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
import tiktoken
from dotenv import load_dotenv
import os
import openai
import streamlit as st
from pathlib import Path

# â—¼ Carrega ambiente e configuraÃ§Ãµes
try:
    api_key = st.secrets["OPENAI_API_KEY"]
    id_bibliografico = st.secrets["ASSISTENTE_BIBLIOGRAFICO"]
except ImportError:
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    id_bibliografico = os.getenv("ASSISTENTE_BIBLIOGRAFICO")

openai.api_key = api_key
ASSISTANT_BIBLIO = id_bibliografico

PASTA_SAIDA = "saida"
VALOR_INPUT = 0.01
VALOR_OUTPUT = 0.03
COTACAO_DOLAR = 5.65
ENCODER = tiktoken.encoding_for_model("gpt-4")
MAX_WORKERS = 10
TIMEOUT = 60

# â”€â”€ UtilitÃ¡rios â”€â”€

def contar_tokens(texto):
    return len(ENCODER.encode(texto))

def extrair_completo(resp):
    try:
        original = re.search(r'âŒ\s*Original:\s*"(.*?)"', resp, re.DOTALL)
        corrigido = re.search(r'âœ…\s*Corrigido:\s*"(.*?)"', resp, re.DOTALL)
        comentario = re.search(r'ğŸ“\s*Justificativa:\s*"(.*?)"', resp, re.DOTALL)
        return (
            original.group(1).strip() if original else None,
            corrigido.group(1).strip() if corrigido else None,
            comentario.group(1).strip() if comentario else None
        )
    except:
        return None, None, None

# â”€â”€ ValidaÃ§Ã£o externa â”€â”€

def validar_doi(doi):
    try:
        r = requests.head(f"https://doi.org/{doi}", timeout=5)
        return r.status_code == 200
    except:
        return False

def validar_isbn(texto):
    return bool(re.search(r"\b(?:ISBN[- ]?)?(97[89])?\d{1,5}[- ]?\d+[- ]?\d+[- ]?\d[Xx]\b", texto))

def validar_url(texto):
    urls = re.findall(r'https?://[^\s]+', texto)
    for u in urls:
        try:
            r = requests.head(u, timeout=5)
            if r.status_code == 200:
                return True
        except:
            continue
    return False

# â”€â”€ Chamada OpenAI â”€â”€

def acionar_assistant(prompt: str, assistant_id: str) -> str | None:
    try:
        import warnings
        warnings.filterwarnings("ignore", category=DeprecationWarning)

        inicio = time.time()

        # Cria um thread (ainda necessÃ¡rio na Assistants API atual)
        thread = openai.beta.threads.create()

        # Adiciona mensagem do usuÃ¡rio
        openai.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content=prompt
        )

        # Roda o assistant e espera automaticamente
        run = openai.beta.threads.runs.create_and_poll(
            thread_id=thread.id,
            assistant_id=assistant_id
        )

        # Verifica status
        if run.status != "completed":
            print(f"âŒ Run falhou: status = {run.status}")
            return None

        # Pega a resposta da IA
        mensagens = openai.beta.threads.messages.list(thread_id=thread.id)
        for msg in reversed(mensagens.data):
            if msg.role == "assistant":
                fim = time.time()
                print(f"âœ… Assistant respondeu em {round(fim - inicio, 2)}s")
                return msg.content[0].text.value.strip()

    except Exception as e:
        print(f"âš ï¸ Erro ao acionar assistant: {e}")
        return None


# â”€â”€ RevisÃ£o de parÃ¡grafo bibliogrÃ¡fico â”€â”€

def revisar_biblio(item, parags):
    if item.get("categoria") != "bibliografico":
        return None
    idx = item["index"]
    if idx < 0 or idx >= len(parags):
        return None
    texto = parags[idx].text.strip()
    prompt = (
        "VocÃª Ã© um assistente especializado em revisar e padronizar referÃªncias segundo ABNT NBR 6023:2018."
        " Mantenha todos os placeholders internos entre colchetes intactos."
        " Detecte e corrija datas mal formatadas: normalize sempre para DD/MM/AAAA ou MÃªs/Ano conforme ABNT."
        " Se o ano estiver incompleto (ex.: 200â€“), sinalize como inconsistÃªncia."
        " Responda no formato:\nâŒ Original: \"â€¦\"\nâœ… Corrigido: \"â€¦\"\nğŸ“ Justificativa: \"â€¦\"\n"
        f"Trecho:\n{texto}"
    )
    resp = acionar_assistant(prompt, ASSISTANT_BIBLIO)
    original, corr, just = extrair_completo(resp or "")
    if corr and corr.lower() != texto.lower():
        in_tk = contar_tokens(prompt)
        out_tk = contar_tokens(resp)
        obs = []
        doi = re.search(r"10\.\d{4,9}/[-._;()/:A-Z0-9]+", corr or texto, re.I)
        obs.append("DOI vÃ¡lido" if doi and validar_doi(doi.group(0)) else "Sem DOI vÃ¡lido")
        obs.append("ISBN" if validar_isbn(corr or texto) else "Sem ISBN")
        obs.append("URL acessÃ­vel" if validar_url(corr or texto) else "Sem URL acessÃ­vel")
        comentario = f"{just or '-'} | {' | '.join(obs)}"
        return {"index": idx, "corrigido": corr, "comentario": comentario,
                "tokens_input": in_tk, "tokens_output": out_tk}
    return None

# â–¶ï¸ Pipeline de aplicaÃ§Ã£o

def aplicar(nomes=None, usuario=""):
    # define pastas a processar
    if nomes:
        to_process = nomes
    else:
        to_process = [d for d in os.listdir(PASTA_SAIDA) 
                      if os.path.isdir(os.path.join(PASTA_SAIDA, d))]

    pasta_base = os.path.join(PASTA_SAIDA, usuario)

    for nome in to_process:
        pasta = os.path.join(pasta_base, nome)
        docx_text = os.path.join(pasta, f"{nome}_revisado_texto.docx")
        map_json = os.path.join(pasta, "mapeamento_textual.json")
        if not os.path.exists(docx_text) or not os.path.exists(map_json):
            print(f"â­ï¸ Pulando {nome}: arquivo nÃ£o encontrado")
            continue

        doc = Document(docx_text)
        with open(map_json, 'r', encoding='utf-8') as f:
            mapa = json.load(f)

        # mantÃ©m estrutura completa
        parags = list(doc.paragraphs)
        revisoes, ti_sum, to_sum = [], 0, 0
        start = time.time()

        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            futuros = {executor.submit(revisar_biblio, item, parags): item 
                       for item in mapa}
            for fut in tqdm(as_completed(futuros), total=len(futuros), desc=f"ğŸ” Bibliografia {nome}"):
                res = fut.result(timeout=TIMEOUT)
                if not res: continue
                idx = res['index']
                parags[idx].text = res['corrigido']
                revisoes.append(res)
                ti_sum += res['tokens_input']; to_sum += res['tokens_output']

        # salva doc com revisÃµes
        doc.save(os.path.join(pasta, f"{nome}_revisado_biblio.docx"))

        # atualiza planilha e relatÃ³rio
        plan = os.path.join(pasta, "avaliacao_completa.xlsx")
        wb = load_workbook(plan) if os.path.exists(plan) else Workbook()
        aba = wb["BibliogrÃ¡fica"] if "BibliogrÃ¡fica" in wb.sheetnames else wb.create_sheet("BibliogrÃ¡fica")
        if aba.max_row == 1:
            aba.append(["ParÃ¡grafo","Texto Corrigido","ComentÃ¡rio IA"])
        for rev in revisoes:
            aba.append([rev['index']+1, rev['corrigido'], rev['comentario']])

        # resumo
        resumo = wb["Resumo"] if "Resumo" in wb.sheetnames else wb.create_sheet("Resumo")
        if resumo.max_row == 1:
            resumo.append(["Revisor","Tempo (s)","Tokens In","Tokens Out","USD","BRL"])
        elapsed = round(time.time()-start, 1)
        usd = (ti_sum*VALOR_INPUT + to_sum*VALOR_OUTPUT)/1000
        resumo.append(["BibliogrÃ¡fico", elapsed, ti_sum, to_sum, round(usd,4), round(usd*COTACAO_DOLAR,2)])
        wb.save(plan)

        # relatÃ³rio tÃ©cnico
        rel_path = os.path.join(pasta, f"relatorio_tecnico_{nome}.docx")
        rel = Document(rel_path) if os.path.exists(rel_path) else Document()
        if revisoes:
            rel.add_page_break()
            rel.add_heading("2. VerificaÃ§Ã£o BibliogrÃ¡fica", level=1)
            for rev in revisoes:
                p = rel.add_paragraph(); p.add_run(f"ParÃ¡grafo {rev['index']+1}:").bold = True; p.add_run(" " + rev['comentario'])
        rel.save(rel_path)

if __name__ == "__main__":
    try:
        if len(sys.argv) >= 3:
            entrada = sys.argv[1]
            usuario = sys.argv[2]
            nome = Path(entrada).stem
            aplicar([nome], usuario)
    except Exception:
        print("âŒ Erro na revisÃ£o bibliogrÃ¡fica:\n", traceback.format_exc())
