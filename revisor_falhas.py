# ✅ Revisor de falhas com agrupamento e revisão contextual (sem mapeamento), com verificação de datas e log no terminal
import os
import sys
import time
import re
from pathlib import Path
from openai import OpenAI
import tiktoken
from docx import Document
from openpyxl import load_workbook, Workbook
from tqdm import tqdm
from dotenv import load_dotenv
from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
import openai
import streamlit as st

# --- Configurações iniciais ---
# 🔁 Compatível com .env local e st.secrets no Streamlit Cloud
try:
    api_key = st.secrets["OPENAI_API_KEY"]
except ImportError:
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")

openai.api_key = api_key
PASTA_SAIDA = "saida"
VALOR_INPUT = 0.01
VALOR_OUTPUT = 0.03
COTACAO_DOLAR = 5.65
ENCODER = tiktoken.encoding_for_model("gpt-4")
TIMEOUT_SEC = 90
MAX_WORKERS = 4
MAX_RETRY = 2

# --- Prompt ---
PROMPT_REVISAO = (
    "Você é um revisor técnico e de estilo com foco em textos acadêmicos e científicos.\n"
    "Corrija blocos de texto apenas se houver erros de gramática, ortografia, datas mal formatadas (como \"13/03/23\" ou \"202-\"), clareza, coesão ou lógica textual.\n"
    "Preserve o estilo do autor e a terminologia técnica.\n"
    "Em trechos com data no formato Mês/Ano com o mês por extenso, mantenha o formato. Por exemplo: 'Março/2023'.\n"
    "Se houver citação bibliográfica com datas incorretas ou incompletas, corrija ou sinalize de forma padronizada conforme a norma ABNT.\n"
    "Responda apenas com o texto revisado, sem explicações."
)

# --- Funções auxiliares ---
def contar_tokens(texto):
    return len(ENCODER.encode(texto))

def tentar_revisar(prompt):
    for _ in range(MAX_RETRY):
        try:
            resp = openai.api_key.chat.completions.create(
                model="gpt-4o",
                messages=[{"role": "user", "content": prompt}]
            )
            return resp.choices[0].message.content.strip()
        except Exception as e:
            print("⚠️ Tentativa de revisão falhou:", e)
            time.sleep(1)
    return None

def pular_paragrafo(texto):
    return (
        texto.isupper() and len(texto.split()) <= 12
        or re.search(r'(sum[aá]rio|lista de fig|conte[úu]do|relat[ óo]rio|impacto.*\d{4})', texto.lower())
        or re.match(r'^\d+(\.\d+)*\s+[A-ZÁÉÍÓÚ]', texto)
    )

def agrupar_paragrafos(parags, max_bloco=3):
    blocos = []
    bloco_atual = []
    indices = []
    for i, p in enumerate(parags):
        texto = p.text.strip()
        if not texto or pular_paragrafo(texto):
            continue
        bloco_atual.append(texto)
        indices.append(i)
        if len(bloco_atual) == max_bloco:
            blocos.append(("\n\n".join(bloco_atual), indices.copy()))
            bloco_atual, indices = [], []
    if bloco_atual:
        blocos.append(("\n\n".join(bloco_atual), indices.copy()))
    return blocos

# --- Execução principal ---
def aplicar(nomes, usuario=""):
    to_process = [n for n in nomes if os.path.isdir(os.path.join(PASTA_SAIDA, n))]
    pasta_base = os.path.join(PASTA_SAIDA, usuario)
    for nome in to_process:
        pasta = os.path.join(pasta_base, nome)
        docx_biblio = os.path.join(pasta, f"{nome}_revisado_biblio.docx")
        docx_texto = os.path.join(pasta, f"{nome}_revisado_texto.docx")

        if os.path.exists(docx_biblio):
            path_docx = docx_biblio
        elif os.path.exists(docx_texto):
            path_docx = docx_texto
        else:
            print(f"⚠️ Nenhum documento revisado encontrado para {nome}, pulando.")
            continue

        print(f"📂 Processando: {nome}")
        doc = Document(path_docx)
        parags = [p for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    parags.extend([p for p in cell.paragraphs if p.text.strip()])

        blocos = agrupar_paragrafos(parags, max_bloco=3)
        revisoes, ti_sum, to_sum = [], 0, 0
        inicio = time.time()

        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            futures = {
                executor.submit(tentar_revisar, PROMPT_REVISAO + f"\nTrecho:\n{bloco}"): idxs
                for bloco, idxs in blocos
            }
            for fut in tqdm(as_completed(futures), total=len(futures), desc=f"🔧 Revisando {nome}"):
                try:
                    resp = fut.result(timeout=TIMEOUT_SEC)
                except Exception:
                    print("⚠️ Timeout em um bloco de revisão.")
                    continue
                if resp:
                    idxs = futures[fut]
                    partes = [s.strip() for s in resp.split("\n\n") if s.strip()]
                    for i, idx in enumerate(idxs):
                        if i < len(partes):
                            print(f"✅ Parágrafo {idx+1} revisado.")
                            ti = contar_tokens(PROMPT_REVISAO + f"\nTrecho:\n{parags[idx].text.strip()}")
                            to = contar_tokens(partes[i])
                            parags[idx].text = partes[i]
                            revisoes.append((idx, partes[i]))
                            ti_sum += ti
                            to_sum += to

        doc.save(os.path.join(pasta, f"{nome}_revisado_completo.docx"))
        print(f"📁 Documento salvo: {nome}_revisado_completo.docx")

        # Planilha
        plan = os.path.join(pasta, "avaliacao_completa.xlsx")
        wb = load_workbook(plan) if os.path.exists(plan) else Workbook()
        aba = wb["Falhas"] if "Falhas" in wb.sheetnames else wb.create_sheet("Falhas")
        if aba.max_row == 1:
            aba.append(["Parágrafo", "Texto Corrigido"])
        for idx, texto in revisoes:
            aba.append([idx + 1, texto])

        resumo = wb["Resumo"] if "Resumo" in wb.sheetnames else wb.create_sheet("Resumo")
        if resumo.max_row == 1:
            resumo.append(["Revisor", "Tempo (s)", "Tokens In", "Tokens Out", "USD", "BRL"])
        tempo = round(time.time() - inicio, 1)
        usd = (ti_sum * VALOR_INPUT + to_sum * VALOR_OUTPUT) / 1000
        resumo.append(["Falhas", tempo, ti_sum, to_sum, round(usd, 4), round(usd * COTACAO_DOLAR, 2)])
        wb.save(plan)
        print(f"📊 Tokens totais: {ti_sum} in / {to_sum} out. Tempo: {tempo}s.\n")

if __name__ == "__main__":
    try:
        if len(sys.argv) >= 3:
            entrada = sys.argv[1]
            usuario = sys.argv[2]
            nome = Path(entrada).stem
            aplicar([nome], usuario)
    except Exception as e:
        print(f"❌ Erro na revisão de falhas: {e}")

