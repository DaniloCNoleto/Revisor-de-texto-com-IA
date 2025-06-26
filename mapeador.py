# 📌 Mapeador de Revisão Dossel ajustado (completo, comentado e categorizado)
import os
import re
import json
import tiktoken
import openai
import time
from docx import Document
from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
from openpyxl import Workbook
from dotenv import load_dotenv

# 🔐 Carrega variáveis de ambiente
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY")
id_bibliografico = os.getenv("ASSISTENTE_BIBLIOGRAFICO")
id_textual = os.getenv("ASSISTENTE_REVISOR_TEXTUAL")
id_tecnico = os.getenv("ASSISTENTE_TECNICO")

# 🔧 Configurações de API e modelo
openai.api_key = api_key
ASSISTANT_TEXTUAL = id_textual
ASSISTANT_ABNT = id_bibliografico
ENCODER = tiktoken.encoding_for_model("gpt-4")

# 📁 Pastas de entrada e saída
PASTA_ENTRADA = "entrada"
PASTA_SAIDA = "saida"

# 📜 Prompts base
PROMPT_ABNT = (
    "Você é um especialista em revisão de textos técnicos segundo as normas ABNT.\n"
    "Avalie se o parágrafo abaixo segue corretamente a formatação, estilo e estrutura exigida.\n"
    "Considere elementos como citação direta/indireta, numeração de seções, uso de abreviações, margens, coerência entre citações e referências.\n"
    "Responda apenas 'Sim' se estiver conforme ou 'Não' se não estiver conforme, seguido de uma breve justificativa.\n"
)

PROMPT_AVALIACAO_TEXTO = (
    "Você é um revisor técnico. Classifique o trecho abaixo indicando se há problemas nas seguintes categorias:\n"
    "ortografia, clareza, coerência, lógica, coesão, consistência.\n"
    "Responda com uma lista separada por vírgulas com as categorias encontradas ou 'nenhum' se não houver problemas.\n"
)

# 🧹 Coleta todos os parágrafos com conteúdo significativo do documento

def coletar_paragrafos(doc):
    paragrafos = [p.text.strip() for p in doc.paragraphs if p.text.strip() and len(p.text.strip()) > 15]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip() and len(p.text.strip()) > 15:
                        paragrafos.append(p.text.strip())
    return paragrafos

# 🧠 Categorização final para execução condicional dos revisores

def classificar_categoria(tipos):
    tipos = [t.lower() for t in tipos]
    if any(t in tipos for t in ["ortografia", "gramática", "sintaxe"]):
        return "textual"
    if any(t in tipos for t in ["referência", "citação", "abnt"]):
        return "bibliografico"
    if any(t in tipos for t in ["clareza", "coerência", "lógica", "coesão", "consistência"]):
        return "estrutura"
    return "nenhuma"

# 🔠 Detecta se trecho é título ou sumário

def detectar_titulo(texto):
    if texto.isupper() and len(texto) < 150:
        return True
    if re.match(r"^\d+(\.\d+)*\s", texto):
        return True
    return False

def detectar_sumario(texto):
    if "sumário" in texto.lower():
        return True
    if re.match(r"^\d+(\.\d+)*\s+.+\s+\.{2,}\s*\d{1,3}$", texto):
        return True
    return False

# 🔢 Conta tokens de um texto

def contar_tokens(texto):
    return len(ENCODER.encode(texto))

# 🤖 Avalia um trecho com um prompt e retorna resposta + metadados

def avaliar_trecho(prompt_base, trecho, assistant_id, retries=3):
    prompt = prompt_base + f"\nTrecho:\n{trecho}"
    tokens_prompt = contar_tokens(prompt)

    for tentativa in range(retries):
        try:
            import warnings
            warnings.filterwarnings("ignore", category=DeprecationWarning)

            inicio = time.time()

            # Cria thread e envia mensagem
            thread = openai.beta.threads.create()
            openai.beta.threads.messages.create(
                thread_id=thread.id,
                role="user",
                content=prompt
            )

            # Roda assistant
            run = openai.beta.threads.runs.create_and_poll(
                thread_id=thread.id,
                assistant_id=assistant_id
            )

            # Se completou com sucesso
            if run.status == "completed":
                fim = time.time()
                msgs = openai.beta.threads.messages.list(thread_id=thread.id)
                for msg in reversed(msgs.data):
                    if msg.role == "assistant":
                        resposta = msg.content[0].text.value.strip()
                        tokens_resposta = contar_tokens(resposta)
                        print(f"✅ Revisão concluída em {round(fim - inicio, 2)}s")
                        return resposta, tokens_prompt, tokens_resposta, round(fim - inicio, 2)
            else:
                print(f"❌ Run falhou: status = {run.status}")

        except Exception as e:
            print(f"⚠️ Tentativa {tentativa+1} falhou: {e}")
            time.sleep(1)

    return "[FALHA]", tokens_prompt, 0, 0



# 🧩 Mapeia um parágrafo individual

def mapear_paragrafo(i, trecho):
    resultado = {
        "paragrafo": i + 1,
        "index": i,
        "tipo": [],
        "falhas": [],
        "categoria": "nenhuma"
    }

    if detectar_sumario(trecho):
        resultado["categoria"] = "sumario"
        return resultado, 0, 0, 0

    if detectar_titulo(trecho):
        resultado["tipo"].append("titulo")
        if len(trecho) > 120:
            resultado["falhas"].append("titulo_longo")

    if re.search(r"(estudo|relatório|avaliação).*impacto.*", trecho.lower()) and re.search(r"(janeiro|fevereiro|março|abril|maio|junho|julho|agosto|setembro|outubro|novembro|dezembro)/\d{4}", trecho.lower()):
        resultado["tipo"].append("capa")

    r_texto, tokens_in1, tokens_out1, tempo1 = avaliar_trecho(PROMPT_AVALIACAO_TEXTO, trecho, ASSISTANT_TEXTUAL)
    if r_texto == "[FALHA]":
        resultado["falhas"].append("avaliacao_texto")
    elif r_texto.strip().lower() != "nenhum":
        resultado["tipo"].extend([x.strip() for x in r_texto.split(",") if x.strip()])
    else:
        resultado["tipo"].append("nenhum")

    r_abnt, tokens_in2, tokens_out2, tempo2 = avaliar_trecho(PROMPT_ABNT, trecho, ASSISTANT_ABNT)
    if r_abnt == "[FALHA]":
        resultado["falhas"].append("avaliacao_abnt")
    elif "não" in r_abnt.lower():
        resultado["tipo"].append("abnt")

    if re.search(r"\b\w+\s*\(\d{4}\)", trecho):
        resultado["tipo"].append("referência")

    resultado["categoria"] = classificar_categoria(resultado["tipo"])
    return resultado, tokens_in1 + tokens_in2, tokens_out1 + tokens_out2, tempo1 + tempo2

# 📄 Processa um documento inteiro .docx

def mapear_documento(nome_arquivo):
    doc_path = os.path.join(PASTA_ENTRADA, nome_arquivo)
    doc = Document(doc_path)
    paragrafos = coletar_paragrafos(doc)
    print(f"\U0001f50d Mapeando {nome_arquivo}... Total: {len(paragrafos)} parágrafos")

    resultados = []
    registros_tokens = []
    nome_saida = os.path.splitext(nome_arquivo)[0]
    os.makedirs(os.path.join(PASTA_SAIDA, nome_saida), exist_ok=True)
    path_json = os.path.join(PASTA_SAIDA, nome_saida, "mapeamento_textual.json")
    path_tokens = os.path.join(PASTA_SAIDA, nome_saida, "mapeamento_tokens.xlsx")

    with ThreadPoolExecutor(max_workers=max(10, os.cpu_count() or 1)) as executor:
        futuros = {executor.submit(mapear_paragrafo, i, p): i for i, p in enumerate(paragrafos)}
        progresso = tqdm(total=len(futuros), desc=f"🔎 Avaliando {nome_arquivo}")
        for fut in as_completed(futuros):
            try:
                resultado, tok_in, tok_out, tempo = fut.result(timeout=120)
                resultados.append(resultado)
                registros_tokens.append([resultado["paragrafo"], tok_in, tok_out, tempo])
            except TimeoutError:
                print("❌ Timeout individual atingido. Pulando parágrafo.")
                registros_tokens.append(["TIMEOUT", 0, 0, 90])
            progresso.update(1)
        progresso.close()

    with open(path_json, "w", encoding="utf-8") as f:
        json.dump(resultados, f, indent=2, ensure_ascii=False)
    print(f"✅ Mapeamento salvo em {path_json}")

    wb = Workbook()
    ws = wb.active
    ws.title = "MapaTokens"
    ws.append(["Parágrafo", "Tokens Prompt", "Tokens Resposta", "Tempo (s)"])
    for linha in registros_tokens:
        ws.append(linha)
    wb.save(path_tokens)
    print(f"📊 Planilha de tokens salva em {path_tokens}")

# ▶️ Execução principal

def main():
    arquivos = [f for f in os.listdir(PASTA_ENTRADA) if f.endswith(".docx") and not f.startswith("~$")]
    for nome in arquivos:
        mapear_documento(nome)

if __name__ == "__main__":
    main()
