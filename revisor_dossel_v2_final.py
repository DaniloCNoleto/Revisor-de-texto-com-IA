# 📌 Revisor Textual Dossel ajustado (com timeout, retries, registro de falhas e suporte a nomes dinâmicos)
import os
import sys
import time
import json
import traceback
import re
import difflib
import openai
import tiktoken
from docx import Document
from openpyxl import Workbook, load_workbook
from tqdm import tqdm
from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError
from dotenv import load_dotenv

# Carrega variáveis de ambiente e configurações
defaults = dict(
    OPENAI_API_KEY=None,
    ASSISTENTE_REVISOR_TEXTUAL=None,
)
load_dotenv()
api_key = os.getenv("OPENAI_API_KEY") or defaults["OPENAI_API_KEY"]
id_textual = os.getenv("ASSISTENTE_REVISOR_TEXTUAL") or defaults["ASSISTENTE_REVISOR_TEXTUAL"]
openai.api_key = api_key
ASSISTANT_TEXTUAL = id_textual

# Configurações gerais
PASTA_SAIDA = "saida"
PASTA_ENTRADA = "entrada"
# Custos
VALOR_INPUT = 0.005
VALOR_OUTPUT = 0.015
COTACAO_DOLAR = 5.65
ENCODER = tiktoken.encoding_for_model("gpt-4")
# Timeouts e tentativas
timeout_seconds = 180
max_retries = 3
result_wait = 10

# Prompts base
PROMPT_REVISAO = (
    "Você é um revisor de textos técnicos com padrão de excelência.\n"
    "Corrija apenas se houver erro. Mantenha o estilo de escrita do autor.\n"
    "Busque padronizar termos técnicos e melhorar a clareza sem alteração no estilo de escrita.\n"
    "Busque manter a formatação original, como negrito e itálico e estrutura do parágrafo, quebras de linha e espaçamento, manter a voz ativa e evitar passivas desnecessárias, manter a terminologia técnica e científica adequada, manter a coerência e coesão do texto, manter a formalidade e objetividade do texto técnico,  manter a clareza e fluidez do texto, manter a precisão e exatidão das informações,  manter a lógica e a argumentação do texto e outros termos e abreviações de acordo com a ABNT, manter a concisão, evitar redundâncias\n"
    "Responda no formato:\n❌ Original: \"...\"\n✅ Corrigido: \"...\"\n📜 Comentário: \"...\"\n"
)
PROMPT_FORMATACAO_FIXA = (
    "Você é um revisor gramatical.\n"
    "Corrija apenas erros de gramática e ortografia. Mantenha estilo de escrita e concisão.\n"
    "Em trechos com data no formato Mês/Ano com o mês por extenso, mantenha o formato. Por exemplo: 'Março/2023'.\n"
    "Responda no formato:\n❌ Original: \"...\"\n✅ Corrigido: \"...\"\n📜 Comentário: \"...\"\n"
)

# Helpers

def workers_dinamicos(minimos=10):
    try:
        return max(minimos, os.cpu_count() or 1)
    except:
        return minimos


def contar_tokens(txt: str) -> int:
    return len(ENCODER.encode(txt))


def similaridade(a: str, b: str) -> float:
    return difflib.SequenceMatcher(None, a.lower(), b.lower()).ratio()


def extrair_completo(resp: str):
    try:
        o = re.search(r'❌\s*Original:\s*["“](.*?)["”]', resp, re.DOTALL)
        c = re.search(r'✅\s*Corrigido:\s*["“](.*?)["”]', resp, re.DOTALL)
        m = re.search(r'📜\s*Comentário:\s*["“](.*?)["”]', resp, re.DOTALL)
        return (
            o.group(1).strip() if o else None,
            c.group(1).strip() if c else None,
            m.group(1).strip() if m else None,
        )
    except:
        return None, None, None

# Executa chamada à API com timeout

def acionar_assistant(prompt: str, assistant_id: str) -> str | None:
    try:
        thread = openai.beta.threads.create()
        openai.beta.threads.messages.create(thread_id=thread.id, role="user", content=prompt)
        run = openai.beta.threads.runs.create(thread_id=thread.id, assistant_id=assistant_id)
        inicio = time.time()
        while True:
            if time.time() - inicio > timeout_seconds:
                print(f"⏲️ Timeout de {timeout_seconds}s (chars={len(prompt)})")
                return None
            run = openai.beta.threads.runs.retrieve(thread_id=thread.id, run_id=run.id)
            if run.status == "completed":
                break
            if run.status in ["failed","cancelled"]:
                return None
            time.sleep(2)
        msgs = openai.beta.threads.messages.list(thread_id=thread.id)
        for msg in reversed(msgs.data):
            if msg.role == "assistant":
                return msg.content[0].text.value
    except:
        return None

# Revisão de parágrafo

def revisar_paragrafo(item: dict, parags: list) -> dict | None:
    if item.get("categoria") != "textual":
        return None
    idx = item["index"]
    if idx < 0 or idx >= len(parags):
        return None
    texto = parags[idx].text.strip()
    tipos = item.get("tipo", []).copy()
    ajustes = []
    # define base do prompt
    prompt_base = PROMPT_REVISAO
    # correção leve para capas detectadas por regex (ex.: 'estudo impacto' + 'mês/ano')
    if re.search(r"(estudo|relatório|avaliação).*impacto.*", texto.lower()) and \
       re.search(r"(janeiro|fevereiro|março|abril|maio|junho|julho|agosto|setembro|outubro|novembro|dezembro)/\d{4}", texto.lower()):
        tipos.append("capa")
        ajustes.append("manter capa")
        prompt_base = PROMPT_FORMATACAO_FIXA
    # correção leve para títulos curtos em maiúsculas
    elif len(texto.split()) <= 12 and texto.isupper():
        tipos.append("título curto em maiúsculas")
        ajustes.append("manter título curto e visual")
        prompt_base = PROMPT_FORMATACAO_FIXA
    # monta prompt final
    prompt = (f"{prompt_base}\nObjetivo: {', '.join(tipos)}\n"
              f"Trecho:\n{texto}")
    tokens_in = contar_tokens(prompt)

    # tenta várias vezes
    for _ in range(max_retries):
        resp = acionar_assistant(prompt, ASSISTANT_TEXTUAL)
        if not resp:
            continue
        original, corr, com = extrair_completo(resp)
        # fallback quando similaridade baixa
        if original and similaridade(original, texto) < 0.75:
            prompt_base = PROMPT_FORMATACAO_FIXA
            prompt = f"{prompt_base}\nTrecho:\n{texto}"
            tokens_in = contar_tokens(prompt)
            resp = acionar_assistant(prompt, ASSISTANT_TEXTUAL)
            if not resp:
                continue
            original, corr, com = extrair_completo(resp)
        if corr and corr.lower() != texto.lower():
            tokens_out = contar_tokens(resp)
            return {
                "index": idx,
                "original": texto,
                "corrigido": corr,
                "comentario": com or resp,
                "tokens_input": tokens_in,
                "tokens_output": tokens_out,
                "ajustes": ajustes
            }
    return None

# Função principal

def aplicar(nomes: list[str] | None = None):
    to_process = nomes or [d for d in os.listdir(PASTA_SAIDA) 
                            if os.path.isdir(os.path.join(PASTA_SAIDA, d))]

    for nome in to_process:
        pasta = os.path.join(PASTA_SAIDA, nome)
        docx_path = os.path.join(PASTA_ENTRADA, nome + ".docx")
        json_map = os.path.join(pasta, "mapeamento_textual.json")
        if not os.path.exists(docx_path) or not os.path.exists(json_map):
            print(f"⏭️ Pulando {nome}")
            continue

        doc = Document(docx_path)
        with open(json_map, "r", encoding="utf-8") as f:
            mapa = json.load(f)

        # extrai parágrafos >15 chars e de tabelas
        parags = [p for p in doc.paragraphs 
                  if p.text.strip() and len(p.text.strip()) > 15]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip() and len(p.text.strip()) > 15:
                            parags.append(p)

        revisoes, failures = [], []
        tokens_in = tokens_out = 0
        start_all = time.time()

        with ThreadPoolExecutor(max_workers=workers_dinamicos()) as executor:
            futures = {executor.submit(revisar_paragrafo, it, parags): it for it in mapa}
            for fut in tqdm(as_completed(futures), total=len(futures),
                             desc=f"🔎 {nome}"):
                item = futures[fut]
                try:
                    texto_item = parags[item["index"]].text
                except:
                    texto_item = ""
                try:
                    res = fut.result(timeout=result_wait)
                except TimeoutError:
                    failures.append({"index": item["index"], "texto": texto_item})
                    continue
                if not res:
                    failures.append({"index": item["index"], "texto": texto_item})
                    continue
                idx = res["index"]
                # preserva prefixo numérico somente para títulos
                if any('título' in t.lower() for t in item.get('tipo', [])):
                    full = parags[idx].text
                    m = re.match(r"^(\d+(?:\.\d+)*\s+)", full)
                    prefix = m.group(1) if m else ""
                    parags[idx].text = prefix + res["corrigido"]
                else:
                    parags[idx].text = res["corrigido"]
                revisoes.append(res)
                tokens_in += res["tokens_input"]
                tokens_out += res["tokens_output"]

        # salva docx revisado e falhas
        doc.save(os.path.join(pasta, f"{nome}_revisado_texto.docx"))
        if failures:
            with open(os.path.join(pasta, "falhas_textual.json"), "w", encoding="utf-8") as jf:
                json.dump(failures, jf, ensure_ascii=False, indent=2)

        # atualiza planilha e relatório
        plan_path = os.path.join(pasta, "avaliacao_completa.xlsx")
        wb = load_workbook(plan_path) if os.path.exists(plan_path) else Workbook()
        if "Texto" not in wb.sheetnames:
            aba = wb.create_sheet("Texto")
            aba.append(["Parágrafo","Tipo","Texto Corrigido"])
        else:
            aba = wb["Texto"]
        for rev in revisoes:
            aba.append([rev["index"]+1, "Textual", rev["corrigido"]])
        if "Resumo" not in wb.sheetnames:
            resumo = wb.create_sheet("Resumo")
            resumo.append(["Revisor","Tempo (s)","Tokens In","Tokens Out","USD","BRL"])
        else:
            resumo = wb["Resumo"]
        tempo = round(time.time()-start_all, 1)
        usd = (tokens_in*VALOR_INPUT + tokens_out*VALOR_OUTPUT)/1000
        resumo.append(["Textual",tempo,tokens_in,tokens_out,round(usd,4),round(usd*COTACAO_DOLAR,2)])
        wb.save(plan_path)

        # gera relatório técnico
        rel_path = os.path.join(pasta, f"relatorio_tecnico_{nome}.docx")
        rel = Document(rel_path) if os.path.exists(rel_path) else Document()
        if revisoes:
            rel.add_page_break()
            rel.add_heading("1. Revisão Textual Técnica", level=1)
            for rev in revisoes:
                par = rel.add_paragraph()
                par.add_run(f"Parágrafo {rev['index']+1}: ").bold = True
                par.add_run(rev.get("comentario",""))
                if rev.get("ajustes"):
                    rel.add_paragraph("Observações adicionais: "+", ".join(rev["ajustes"]),
                                     style="Intense Quote")
        rel.save(rel_path)
        print(f"✅ Revisão aplicada: {nome}")

if __name__ == "__main__":
    try:
        aplicar(sys.argv[1:] if len(sys.argv)>1 else None)
    except Exception:
        print("❌ Erro:", traceback.format_exc())
        sys.exit(1)
