import unittest
import os
import json
import time
from unittest.mock import patch
from docx import Document
from pathlib import Path

# Importações com os nomes reais dos módulos
from revisor_dossel_v2_final import contar_tokens, acionar_assistant
from verificador_bibliografico_final import acionar_assistant_bibliografico
from revisor_falhas import acionar_assistant as acionar_assistant_falhas
from track_changes_final import carregar_textos, comparar_paragrafos

ENTRADA = "entrada"
SAIDA = "saida"

class TestRevisorDossel(unittest.TestCase):

    def setUp(self):
        os.makedirs(ENTRADA, exist_ok=True)
        self.doc_teste = os.path.join(ENTRADA, "teste_unitario.docx")
        self.doc_revisado = os.path.join(ENTRADA, "teste_unitario_revisado.docx")
        self.doc = Document()
        self.doc.add_paragraph("Este é um parágrafo de teste.")
        self.doc.save(self.doc_teste)
        self.doc2 = Document()
        self.doc2.add_paragraph("Este é um parágrafo revisado.")
        self.doc2.save(self.doc_revisado)

    def tearDown(self):
        for f in [self.doc_teste, self.doc_revisado]:
            if os.path.exists(f):
                os.remove(f)

    def test_contar_tokens(self):
        texto = "Teste de contagem de tokens."
        tokens = contar_tokens(texto)
        self.assertIsInstance(tokens, int)
        self.assertGreater(tokens, 0)

    def test_extrair_paragrafos_docx(self):
        doc = Document(self.doc_teste)
        paragrafos = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        self.assertIn("Este é um parágrafo de teste.", paragrafos)

    def test_mock_assistant_textual(self):
        with patch("revisor_dossel_v2_final.acionar_assistant", return_value='Corrigido: "Texto revisado."\nComentário: "Melhoria aplicada."') as mock_func:
            resposta = mock_func("Teste de mock", "asst_fake")
            self.assertIn("Texto revisado", resposta)

    def test_mock_assistant_biblio(self):
        with patch("verificador_bibliografico_final.acionar_assistant_bibliografico", return_value='✅ Corrigido: "Texto bibliográfico."\n📝 Justificativa: "Formatado conforme ABNT."') as mock_func:
            resposta = mock_func("Mock biblio", "asst_biblio")
            self.assertIn("Texto bibliográfico", resposta)

    def test_mock_assistant_falhas(self):
        with patch("revisor_falhas.acionar_assistant", return_value='Corrigido: "Texto corrigido."\nComentário: "Erro gramatical ajustado."') as mock_func:
            resposta = mock_func("Teste falha", "asst_falha")
            self.assertIn("Texto corrigido", resposta)

    def test_track_changes_comparador(self):
        texto1 = carregar_textos(self.doc_teste)
        texto2 = carregar_textos(self.doc_revisado)
        alterados = comparar_paragrafos(texto1, texto2)
        self.assertEqual(len(alterados), 1)
        self.assertEqual(alterados[0][0], 0)
        self.assertIn("teste", alterados[0][1])
        self.assertIn("revisado", alterados[0][2])

if __name__ == '__main__':
    unittest.main()